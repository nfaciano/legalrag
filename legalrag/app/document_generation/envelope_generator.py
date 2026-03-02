"""
Envelope generator for creating print-ready #10 envelope documents.
Return address in absolute-position frame, recipient as centered body text.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, Optional
import logging

logger = logging.getLogger(__name__)


class EnvelopeGenerator:
    """Generates print-ready envelope documents in #10 business envelope format."""

    def generate_envelope(
        self,
        output_path: Path,
        return_address: Dict[str, str],
        recipient_address: Dict[str, str]
    ) -> Path:
        """Generate a print-ready envelope document."""
        try:
            doc = Document()

            # Zero out default style spacing
            style = doc.styles['Normal']
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.0
            style.font.name = 'Arial'
            style.font.size = Pt(12)

            # Page setup: #10 envelope 9.5" x 4.125"
            section = doc.sections[0]
            section.page_width = Inches(9.5)
            section.page_height = Inches(4.125)
            section.top_margin = Inches(0.25)
            section.bottom_margin = Inches(0.25)
            section.left_margin = Inches(0.25)
            section.right_margin = Inches(0.25)

            # Remove default empty paragraph
            for p in doc.paragraphs:
                p._element.getparent().remove(p._element)

            # --- Return address: absolute-position frame (top-left, out of flow) ---
            return_lines = [v for v in [
                return_address.get('name', ''),
                return_address.get('line1', ''),
                return_address.get('line2', ''),
                return_address.get('city_state_zip', ''),
            ] if v.strip()]

            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            framePr = OxmlElement('w:framePr')
            framePr.set(qn('w:w'), str(int(3.5 * 1440)))
            framePr.set(qn('w:h'), str(int(1.0 * 1440)))
            framePr.set(qn('w:x'), str(int(0.5 * 1440)))
            framePr.set(qn('w:y'), str(int(0.35 * 1440)))
            framePr.set(qn('w:hAnchor'), 'page')
            framePr.set(qn('w:vAnchor'), 'page')
            framePr.set(qn('w:wrap'), 'notBeside')
            pPr.append(framePr)

            for i, line in enumerate(return_lines):
                if i > 0:
                    p.add_run().add_break()
                run = p.add_run(line)
                run.font.name = 'Arial'
                run.font.size = Pt(9)

            # Push recipient down to vertical center
            doc.add_paragraph()
            doc.add_paragraph()

            # --- Recipient address: centered body text ---
            recip_lines = [v.upper() for v in [
                recipient_address.get('name', ''),
                recipient_address.get('line1', ''),
                recipient_address.get('line2', ''),
                recipient_address.get('city_state_zip', ''),
            ] if v.strip()]

            rp = doc.add_paragraph()
            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, line in enumerate(recip_lines):
                if i > 0:
                    rp.add_run().add_break()
                run = rp.add_run(line)
                run.font.name = 'Arial'
                run.font.size = Pt(12)

            doc.save(str(output_path))
            logger.info(f"Generated envelope: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error generating envelope: {e}", exc_info=True)
            raise ValueError(f"Failed to generate envelope: {str(e)}")


_envelope_generator: Optional[EnvelopeGenerator] = None


def get_envelope_generator() -> EnvelopeGenerator:
    global _envelope_generator
    if _envelope_generator is None:
        _envelope_generator = EnvelopeGenerator()
    return _envelope_generator

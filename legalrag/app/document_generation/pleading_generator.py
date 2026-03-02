"""
Court pleading document generator using python-docx.
Produces .docx files matching federal court filing format:
caption with colon-aligned parties, centered titles, signature blocks, certification page.

Body text supports markdown-lite formatting:
  ## HEADING        → centered, bold, underlined section heading
  **text**          → bold inline
  __text__          → underlined inline
  1. / a.           → numbered list (hanging indent)
  -                 → bullet list (hanging indent)
  (everything else) → regular paragraph with 0.5" first-line indent
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
import logging

logger = logging.getLogger(__name__)

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)

# Certification templates
FILING_TEMPLATES = {
    "ecf": "I hereby certify that I have electronically filed the within on this {date} using the CM/ECF system",
    "tyler": "I hereby certify that I have filed the within on this {date} via the Tyler electronic filing system",
    "mail_clerk": "I hereby certify that I have filed the within on this {date} by mailing the same to the Clerk of Court",
    "hand_clerk": "I hereby certify that I have filed the within on this {date} by hand delivery to the Clerk of Court",
}

SERVICE_TEMPLATES = {
    "ecf_auto": ", which will send notification of such filing to all counsel of record.",
    "email": ", and that I have served a copy upon counsel of record via email at:",
    "mail": ", and that I have caused a copy to be mailed, first-class, postage prepaid, to:",
    "hand": ", and that I have served a copy by hand delivery upon:",
}

# Inline formatting regex: matches **bold** and __underline__
_INLINE_RE = re.compile(r'(\*\*(.+?)\*\*|__(.+?)__)')

# Numbered list pattern: 1. , 2. , a. , b. , i. , ii. , etc.
_NUMBERED_RE = re.compile(r'^(\d+|[a-zA-Z]{1,4})\.\s+(.+)')


def _set_run_font(run, name=FONT_NAME, size=FONT_SIZE, bold=False, underline=False):
    """Apply standard font formatting to a run."""
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.underline = underline


def _parse_inline_runs(text: str) -> List[Tuple[str, bool, bool]]:
    """Parse markdown-lite inline formatting into (text, bold, underline) tuples."""
    runs = []
    last_end = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > last_end:
            runs.append((text[last_end:match.start()], False, False))
        if match.group(2) is not None:
            runs.append((match.group(2), True, False))
        elif match.group(3) is not None:
            runs.append((match.group(3), False, True))
        last_end = match.end()
    if last_end < len(text):
        runs.append((text[last_end:], False, False))
    if not runs:
        runs.append((text, False, False))
    return runs


def _add_paragraph(doc, text="", alignment=None, bold=False, underline=False,
                   font_size=FONT_SIZE, space_before=Pt(0), space_after=Pt(0),
                   first_line_indent=None, left_indent=None, line_spacing=None):
    """Add a formatted paragraph to the document (single run, no inline formatting)."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if line_spacing is not None:
        pf.line_spacing = line_spacing

    if text:
        run = p.add_run(text)
        _set_run_font(run, size=font_size, bold=bold, underline=underline)
    return p


def _add_rich_paragraph(doc, text, alignment=None, base_bold=False, base_underline=False,
                        font_size=FONT_SIZE, space_before=Pt(0), space_after=Pt(0),
                        first_line_indent=None, left_indent=None, line_spacing=None):
    """Add a paragraph with inline bold/underline formatting parsed from markdown-lite."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if line_spacing is not None:
        pf.line_spacing = line_spacing

    for run_text, is_bold, is_underline in _parse_inline_runs(text):
        run = p.add_run(run_text)
        _set_run_font(
            run, size=font_size,
            bold=base_bold or is_bold,
            underline=base_underline or is_underline,
        )
    return p


def _remove_table_borders(table):
    """Remove all visible borders from a table (make it invisible)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


class PleadingGenerator:
    """Generates court pleading documents in federal court format."""

    def generate_pleading(
        self,
        output_path: Path,
        case_data: Dict[str, Any],
        attorney_info: Dict[str, str],
        document_title: str,
        body_paragraphs: List[str],
        representing_party: str,
        attorney_capacity: str = "By his Attorney,",
        include_certification: bool = True,
        certification_date: Optional[str] = None,
        service_list: Optional[List[str]] = None,
        body_text: Optional[str] = None,
        filing_method: str = "ecf",
        service_method: str = "ecf_auto",
    ) -> Path:
        """Generate a complete court pleading document."""
        doc = Document()
        section = doc.sections[0]

        # Page setup: Letter size, 1-inch margins
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # 1. File reference line
        file_ref = case_data.get("file_reference", "")
        initials = attorney_info.get("initials", "")
        if file_ref or initials:
            ref_text = f"{initials} {file_ref}".strip()
            _add_paragraph(doc, ref_text, font_size=Pt(8), space_after=Pt(12))

        # 2. Caption table (includes court name in right column)
        self._build_caption_table(doc, case_data)

        # 4. Document title
        _add_paragraph(doc, space_before=Pt(18))  # spacer
        _add_paragraph(
            doc, document_title,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True, underline=True,
            space_after=Pt(18),
        )

        # 5. Body content
        if body_text:
            self._add_rich_body(doc, body_text)
        elif body_paragraphs:
            for para_text in body_paragraphs:
                _add_rich_paragraph(
                    doc, para_text,
                    first_line_indent=Inches(0.5),
                    space_after=Pt(6),
                    line_spacing=1.15,
                )

        # 6. Signature block
        _add_paragraph(doc, space_before=Pt(18))  # spacer
        self._add_signature_block(
            doc, representing_party, attorney_capacity, attorney_info
        )

        # 7. Certification page
        if include_certification:
            doc.add_page_break()
            self._add_certification_page(
                doc, attorney_info, certification_date, service_list,
                filing_method=filing_method,
                service_method=service_method,
            )

        doc.save(str(output_path))
        logger.info(f"Generated pleading: {output_path}")
        return output_path

    def _build_caption_table(self, doc: Document, case_data: Dict[str, Any]):
        """Build the caption block as a 3-column invisible-border table.
        Court name is split: state+location top-left, court type top-right.
        Case number lines up with VS row. Colons form a vertical line.
        """
        plaintiffs = case_data["plaintiff_names"]
        defendants = case_data["defendant_names"]
        case_number = case_data["case_number"]
        court_name = case_data.get("court_name", "")
        court_location = case_data.get("court_location", "")
        vs_label = case_data.get("plaintiff_label", "VS.")

        # Split court_name into lines: e.g. "STATE OF RHODE ISLAND\nSUPERIOR COURT"
        court_parts = [l.strip().upper() for l in court_name.split("\n") if l.strip()]
        # First line(s) = state/jurisdiction (left), last line = court type (right)
        if len(court_parts) >= 2:
            court_left_line = court_parts[0]   # e.g. "STATE OF RHODE ISLAND"
            court_right_line = court_parts[-1]  # e.g. "SUPERIOR COURT"
        elif len(court_parts) == 1:
            court_left_line = court_parts[0]
            court_right_line = ""
        else:
            court_left_line = ""
            court_right_line = ""

        location_line = court_location.strip().upper() if court_location else ""

        # Build rows: court header, then blank, then parties with VS + case number
        rows_data = []

        SEP = ")"

        # Row 0: state name (left) | court type (right)
        rows_data.append((court_left_line, SEP, court_right_line))
        # Row 1: location (left) | blank (right)
        rows_data.append((location_line, SEP, ""))
        # Row 2: blank separator
        rows_data.append(("", SEP, ""))

        # Plaintiffs
        for i, name in enumerate(plaintiffs):
            if i > 0 and not name.lower().startswith("and "):
                name = f"and {name}"
            rows_data.append((name, SEP, ""))

        # Blank before VS
        rows_data.append(("", SEP, ""))
        # VS row — case number on right
        vs_row_idx = len(rows_data)
        rows_data.append((f"    {vs_label}", SEP, f"C.A. NO. {case_number}"))
        # Blank after VS
        rows_data.append(("", SEP, ""))

        # Defendants
        for name in defendants:
            rows_data.append((name, SEP, ""))

        # Pad to minimum 9 rows
        while len(rows_data) < 9:
            rows_data.append(("", ":", ""))

        table = doc.add_table(rows=len(rows_data), cols=3)
        _remove_table_borders(table)
        tbl = table._tbl

        # --- Fixed column widths in twips (1 inch = 1440 twips) ---
        # Left parties | colon | right (court name + case number, left-aligned so they share a margin)
        # 3.25 + 0.25 + 3.0 = 6.5 inches (matches usable page width)
        COL_WIDTHS_TWIPS = [int(3.25 * 1440), int(0.25 * 1440), int(3.0 * 1440)]
        # Keep Inches() versions for cell.width (python-docx API expects EMU)
        COL_WIDTHS_EMU = [Inches(3.25), Inches(0.25), Inches(3.0)]
        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is None:
            tblGrid = OxmlElement('w:tblGrid')
            tbl.insert(1, tblGrid)
        for existing in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(existing)
        for w in COL_WIDTHS_TWIPS:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(w))
            tblGrid.append(gridCol)

        # Fixed table layout
        tblPr = tbl.tblPr
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        # Zero out default table cell margins so rows pack tightly
        tblCellMar = OxmlElement('w:tblCellMar')
        for side in ['top', 'bottom']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'), '0')
            el.set(qn('w:type'), 'dxa')
            tblCellMar.append(el)
        tblPr.append(tblCellMar)

        # Exact row height in twips (1pt = 20 twips). 14pt = 280 twips.
        ROW_HEIGHT_TWIPS = 14 * 20  # 280

        for row_idx, (left_text, colon, right_text) in enumerate(rows_data):
            row = table.rows[row_idx]

            # Set exact row height via XML
            trPr = row._tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                row._tr.insert(0, trPr)
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(ROW_HEIGHT_TWIPS))
            trHeight.set(qn('w:hRule'), 'exact')
            trPr.append(trHeight)

            for col_idx, (text, width) in enumerate(
                [(left_text, COL_WIDTHS_EMU[0]), (colon, COL_WIDTHS_EMU[1]), (right_text, COL_WIDTHS_EMU[2])]
            ):
                cell = row.cells[col_idx]
                cell.width = width
                p = cell.paragraphs[0]
                pf = p.paragraph_format
                if col_idx == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif col_idx == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    pf.left_indent = Inches(0.75)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                # Set exact line spacing via XML (twips: 14pt = 280)
                LINE_TWIPS = 14 * 20  # 280
                pPr = p._p.get_or_add_pPr()
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.append(spacing)
                spacing.set(qn('w:line'), str(LINE_TWIPS))
                spacing.set(qn('w:lineRule'), 'exact')
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')

                if text:
                    run = p.add_run(text)
                    _set_run_font(run)

    def _add_rich_body(self, doc: Document, body_text: str):
        """Parse markdown-lite body text and add formatted paragraphs."""
        lines = body_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip blank lines
            if not stripped:
                i += 1
                continue

            # Section heading: ## HEADING TEXT
            if stripped.startswith('## '):
                heading_text = stripped[3:].strip()
                _add_rich_paragraph(
                    doc, heading_text,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    base_bold=True, base_underline=True,
                    space_before=Pt(12), space_after=Pt(6),
                )
                i += 1
                continue

            # Numbered list: 1. , 2. , a. , b. , etc.
            num_match = _NUMBERED_RE.match(stripped)
            if num_match:
                prefix = num_match.group(1) + '.'
                content = num_match.group(2)
                _add_rich_paragraph(
                    doc, f"{prefix}  {content}",
                    left_indent=Inches(0.75),
                    first_line_indent=Inches(-0.25),
                    space_after=Pt(4),
                    line_spacing=1.15,
                )
                i += 1
                continue

            # Bullet list: - text
            if stripped.startswith('- '):
                bullet_text = stripped[2:].strip()
                _add_rich_paragraph(
                    doc, f"\u2022  {bullet_text}",
                    left_indent=Inches(0.75),
                    first_line_indent=Inches(-0.25),
                    space_after=Pt(4),
                    line_spacing=1.15,
                )
                i += 1
                continue

            # Regular paragraph — accumulate consecutive non-blank, non-special lines
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                next_stripped = lines[i].strip()
                if (not next_stripped or
                        next_stripped.startswith('## ') or
                        next_stripped.startswith('- ') or
                        _NUMBERED_RE.match(next_stripped)):
                    break
                para_lines.append(next_stripped)
                i += 1

            full_para = ' '.join(para_lines)
            _add_rich_paragraph(
                doc, full_para,
                first_line_indent=Inches(0.5),
                space_after=Pt(6),
                line_spacing=1.15,
            )

    def _add_signature_block(
        self, doc: Document,
        representing_party: str,
        attorney_capacity: str,
        attorney_info: Dict[str, str],
    ):
        """Add the right-of-center signature block."""
        indent = Inches(3.25)

        _add_paragraph(doc, representing_party, left_indent=indent, space_after=Pt(2))
        _add_paragraph(doc, attorney_capacity, left_indent=indent, space_after=Pt(2))

        for _ in range(3):
            _add_paragraph(doc, left_indent=indent)

        sig_name = attorney_info.get("signature_name", "")
        _add_paragraph(doc, f"/s/ {sig_name}", left_indent=indent, space_after=Pt(2))

        bar = attorney_info.get("bar_number", "")
        name_line = f"{sig_name}, Esq. {bar}".strip() if bar else f"{sig_name}, Esq."
        _add_paragraph(doc, name_line, left_indent=indent, space_after=Pt(1))

        firm = attorney_info.get("firm_name", "")
        if firm:
            _add_paragraph(doc, firm, left_indent=indent, space_after=Pt(1))

        addr1 = attorney_info.get("attorney_address_line1", "")
        if addr1:
            _add_paragraph(doc, addr1, left_indent=indent, space_after=Pt(1))

        addr2 = attorney_info.get("attorney_address_line2", "")
        if addr2:
            _add_paragraph(doc, addr2, left_indent=indent, space_after=Pt(1))

        city = attorney_info.get("attorney_city_state_zip", "")
        if city:
            _add_paragraph(doc, city, left_indent=indent, space_after=Pt(1))

        phone = attorney_info.get("phone", "")
        if phone:
            p = _add_paragraph(doc, left_indent=indent, space_after=Pt(1), font_size=Pt(10))
            run = p.add_run(f"PHONE:  {phone}")
            _set_run_font(run, size=Pt(10))

        fax = attorney_info.get("fax", "")
        if fax:
            p = _add_paragraph(doc, left_indent=indent, space_after=Pt(1), font_size=Pt(10))
            run = p.add_run(f"FAX:  {fax}")
            _set_run_font(run, size=Pt(10))

        email = attorney_info.get("email", "")
        if email:
            p = _add_paragraph(doc, left_indent=indent, space_after=Pt(1), font_size=Pt(10))
            run = p.add_run(f"EMAIL:  {email}")
            _set_run_font(run, size=Pt(10))

    def _add_certification_page(
        self, doc: Document,
        attorney_info: Dict[str, str],
        certification_date: Optional[str],
        service_list: Optional[List[str]],
        filing_method: str = "ecf",
        service_method: str = "ecf_auto",
    ):
        """Add the Certification of Service page with configurable filing/service methods."""
        # Title
        _add_paragraph(doc, space_after=Pt(12))
        _add_paragraph(
            doc, "CERTIFICATION",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True, underline=True,
            space_after=Pt(18),
        )

        # Build certification text from templates
        date_str = certification_date or "___"
        filing_text = FILING_TEMPLATES.get(filing_method, FILING_TEMPLATES["ecf"])
        service_text = SERVICE_TEMPLATES.get(service_method, SERVICE_TEMPLATES["ecf_auto"])
        cert_text = filing_text.format(date=date_str) + service_text

        _add_paragraph(
            doc, cert_text,
            first_line_indent=Inches(0.5),
            space_after=Pt(12),
            line_spacing=1.15,
        )

        # Service list (only for non-ecf_auto methods)
        if service_method != "ecf_auto" and service_list:
            for entry in service_list:
                entry_lines = entry.strip().split("\n")
                for j, line in enumerate(entry_lines):
                    _add_paragraph(
                        doc, line.strip(),
                        space_after=Pt(1) if j < len(entry_lines) - 1 else Pt(12),
                    )

        # Signature
        indent = Inches(3.25)
        _add_paragraph(doc, space_before=Pt(18))
        sig_name = attorney_info.get("signature_name", "")
        _add_paragraph(doc, f"/s/ {sig_name}", left_indent=indent)


# Singleton
_pleading_generator: Optional[PleadingGenerator] = None


def get_pleading_generator() -> PleadingGenerator:
    """Get or create the singleton PleadingGenerator instance."""
    global _pleading_generator
    if _pleading_generator is None:
        _pleading_generator = PleadingGenerator()
    return _pleading_generator

"""
Document builder for assembling and generating documents from templates.
"""

import uuid
import re
import json
from pathlib import Path
from typing import Optional, Dict, Any
from datetime import datetime
import logging

from .template_manager import get_template_manager
from .template_document_generator import get_template_generator

logger = logging.getLogger(__name__)


class DocumentBuilder:
    """Builds complete documents from Word templates."""

    def __init__(self, output_dir: str = "./data/uploads/generated"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.template_manager = get_template_manager()
        self.template_generator = get_template_generator()

    async def generate_from_template(
        self,
        user_id: str,
        template_id: Optional[str],
        fields: Dict[str, str],
        ai_generate_body: bool = False,
        ai_prompt: Optional[str] = None,
        output_format: str = "docx"  # "docx" or "pdf"
    ) -> Dict[str, Any]:
        """
        Generate a document from a Word template.

        Args:
            user_id: User ID from authentication
            template_id: ID of template to use
            fields: Dict of field_name -> value to replace in template
            ai_generate_body: Whether to use AI for body generation
            ai_prompt: Prompt for AI body generation
            output_format: "docx" or "pdf"

        Returns:
            Dict with document_id, filename, and file path
        """
        # Generate body with AI if requested, OR parse all fields if they're empty
        if ai_generate_body and ai_prompt:
            # Check if this is AI Quick mode (all fields empty except defaults)
            is_quick_mode = not fields.get('recipient_name') and not fields.get('body')
            logger.info(f"AI generation requested. Quick mode: {is_quick_mode}, recipient_name: '{fields.get('recipient_name')}', body: '{fields.get('body')[:50] if fields.get('body') else ''}'")

            if is_quick_mode:
                # AI Quick mode - parse all fields from prompt
                logger.info(f"Using AI Quick mode to parse fields from prompt: {ai_prompt[:100]}...")
                parsed_fields = await self._parse_fields_with_ai(ai_prompt)
                logger.info(f"Parsed fields: {parsed_fields}")
                # Merge parsed fields with provided fields
                # User settings (closing, signature_name, initials) take precedence
                for key, value in parsed_fields.items():
                    if key in ['closing', 'signature_name', 'initials']:
                        # Only use AI parsed value if user didn't provide one
                        if not fields.get(key):
                            fields[key] = value
                    elif not fields.get(key) or key in ['body', 'recipient_name', 'recipient_address', 'subject', 'salutation']:
                        # For other fields, use AI parsed value
                        fields[key] = value
            else:
                # Regular mode - just generate body
                logger.info("Using regular mode to generate body only")
                fields['body'] = await self._generate_body_with_ai(ai_prompt)

        # Get template path
        template_path = None
        if template_id:
            template_path = self.template_manager.get_template_path(user_id, template_id)
            if not template_path:
                raise ValueError(f"Template {template_id} not found")

        # If no template, create a basic one on the fly
        if not template_path:
            template_path = self._create_default_template()

        # Generate unique document ID
        doc_id = f"doc_{uuid.uuid4().hex[:12]}"

        # Create user output directory
        user_output_dir = self.output_dir / user_id
        user_output_dir.mkdir(parents=True, exist_ok=True)

        # Generate meaningful filename
        # Use recipient name, subject, or fallback to "letter"
        def sanitize_filename(text: str, max_length: int = 30) -> str:
            """Sanitize text for use in filename"""
            if not text:
                return ""
            # Remove/replace invalid filename characters
            sanitized = re.sub(r'[<>:"/\\|?*]', '', text)
            # Replace spaces and multiple underscores
            sanitized = re.sub(r'\s+', '_', sanitized)
            sanitized = re.sub(r'_+', '_', sanitized)
            # Trim to max length
            return sanitized[:max_length].strip('_')

        # Build filename parts
        parts = []

        # Add recipient name if available
        recipient_name = sanitize_filename(fields.get('recipient_name', ''))
        if recipient_name:
            parts.append(recipient_name)
        else:
            # Fallback to subject if no recipient
            subject = sanitize_filename(fields.get('subject', ''))
            if subject:
                parts.append(subject)
            else:
                parts.append("letter")

        # Add date
        safe_date = fields.get('date', datetime.now().strftime('%Y-%m-%d'))
        safe_date = safe_date.replace("/", "-").replace(" ", "_").replace(",", "")
        parts.append(safe_date)

        # Add doc_id for uniqueness
        parts.append(doc_id)

        filename = "_".join(parts) + ".docx"
        output_path = user_output_dir / filename

        # Generate letter from template (preserves header/footer, builds body)
        self.template_generator.generate_letter_from_template(
            template_path=template_path,
            output_path=output_path,
            fields=fields
        )

        # Convert to PDF if requested
        if output_format == "pdf":
            pdf_filename = filename.replace('.docx', '.pdf')
            pdf_path = user_output_dir / pdf_filename
            final_path = self.template_generator.convert_to_pdf(output_path, pdf_path)

            if final_path == pdf_path:
                # Successfully converted, delete docx
                output_path.unlink()
                output_path = pdf_path
                filename = pdf_filename

        logger.info(f"Generated document {doc_id} for user {user_id}")

        return {
            "document_id": doc_id,
            "filename": filename,
            "file_path": str(output_path),
            "file_size": output_path.stat().st_size,
            "generated_at": datetime.utcnow().isoformat(),
        }

    async def _generate_body_with_ai(self, prompt: str) -> str:
        """
        Generate document body using AI (Groq).

        Args:
            prompt: User prompt for body generation

        Returns:
            Generated body text
        """
        try:
            import os
            from groq import Groq

            api_key = os.environ.get("GROQ_API_KEY")
            if not api_key:
                raise ValueError("GROQ_API_KEY not found in environment")

            client = Groq(api_key=api_key)

            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a professional legal assistant. Generate formal, professional letter body text based on the user's request. Keep it concise and professional."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                max_tokens=1000,
                temperature=0.7,
            )

            body = response.choices[0].message.content.strip()
            return body

        except Exception as e:
            logger.error(f"Error generating body with AI: {e}")
            raise ValueError(f"Failed to generate body with AI: {str(e)}")

    async def _parse_fields_with_ai(self, prompt: str) -> Dict[str, str]:
        """
        Parse all document fields from a natural language prompt using AI.

        Args:
            prompt: Natural language prompt like "Draft a letter to John Smith at 123 Main St..."

        Returns:
            Dict of parsed fields (recipient_name, recipient_address, subject, body, etc.)
        """
        try:
            import os
            from groq import Groq

            api_key = os.environ.get("GROQ_API_KEY")
            if not api_key:
                raise ValueError("GROQ_API_KEY not found in environment")

            client = Groq(api_key=api_key)

            # Create a structured prompt for field extraction
            system_prompt = """You are a legal document assistant. Parse natural language requests into structured JSON for legal letters.

Extract these fields:
- recipient_name: Full name of recipient
- recipient_company: Company name if mentioned
- recipient_address: Full address (use \\n for line breaks)
- subject: Subject/Re: line content
- salutation: How to address (e.g. "Dear Mr. Smith," or "Dear Ms. Jones," - use COMMA, not colon)
- body: Professional letter body content (2-3 paragraphs)
- closing: Letter closing if mentioned (e.g. "Sincerely," or "Best regards,")
- signature_name: Sender's name if mentioned
- initials: Sender's initials if mentioned
- enclosures: List of mentioned enclosures (just the items, don't include "Enc:")

Return ONLY valid JSON. No extra text."""

            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Parse this request:\n\n{prompt}"}
                ],
                max_tokens=1500,
                temperature=0.3,
            )

            response_text = response.choices[0].message.content.strip()
            logger.info(f"AI response for field parsing: {response_text[:200]}...")

            # Try to parse JSON from response (extract from response in case there's extra text)
            json_match = re.search(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                logger.info(f"Extracted JSON: {json_str[:200]}...")
                parsed = json.loads(json_str)
            else:
                # Fallback: try parsing entire response
                logger.info("No JSON pattern found, trying to parse entire response")
                parsed = json.loads(response_text)

            # Ensure all expected fields exist
            fields = {
                'recipient_name': parsed.get('recipient_name', ''),
                'recipient_company': parsed.get('recipient_company', ''),
                'recipient_address': parsed.get('recipient_address', ''),
                'subject': parsed.get('subject', ''),
                'salutation': parsed.get('salutation', 'Dear Sir or Madam,'),
                'body': parsed.get('body', ''),
                'closing': parsed.get('closing', ''),
                'signature_name': parsed.get('signature_name', ''),
                'initials': parsed.get('initials', ''),
                'enclosures': parsed.get('enclosures', ''),
            }

            logger.info(f"AI parsed fields: {list(fields.keys())}")
            return fields

        except Exception as e:
            logger.error(f"Error parsing fields with AI: {e}", exc_info=True)
            logger.warning(f"Falling back to using prompt as body. Original prompt: {prompt[:100]}...")
            # Return fallback with just the body
            return {
                'recipient_name': '',
                'recipient_company': '',
                'recipient_address': '',
                'subject': '',
                'salutation': 'Dear Sir or Madam,',
                'body': prompt,  # Use the original prompt as body
                'enclosures': '',
            }

    def _create_default_template(self) -> Path:
        """
        Create a basic default template if user hasn't uploaded one.

        Returns:
            Path to default template
        """
        from docx import Document

        default_template_path = self.output_dir / "default_template.docx"

        if default_template_path.exists():
            return default_template_path

        # Create basic template
        doc = Document()

        # Add template structure with placeholders
        doc.add_paragraph("{{date}}")
        doc.add_paragraph()
        doc.add_paragraph("{{recipient_name}}")
        doc.add_paragraph("{{recipient_company}}")
        doc.add_paragraph("{{recipient_address}}")
        doc.add_paragraph()
        doc.add_paragraph("Re: {{subject}}")
        doc.add_paragraph()
        doc.add_paragraph("{{salutation}}")
        doc.add_paragraph()
        doc.add_paragraph("{{body}}")
        doc.add_paragraph()
        doc.add_paragraph("{{closing}}")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("{{signature_name}}")
        doc.add_paragraph()
        doc.add_paragraph("{{initials}}")
        doc.add_paragraph("{{enclosures}}")

        doc.save(str(default_template_path))

        logger.info("Created default template")
        return default_template_path


# Singleton instance
_document_builder: Optional[DocumentBuilder] = None


def get_document_builder() -> DocumentBuilder:
    """Get or create the singleton DocumentBuilder instance."""
    global _document_builder
    if _document_builder is None:
        _document_builder = DocumentBuilder()
    return _document_builder

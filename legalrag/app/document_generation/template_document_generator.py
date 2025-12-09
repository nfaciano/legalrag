"""
Template-based document generator using python-docx for placeholder replacement.
"""

from docx import Document
from pathlib import Path
from typing import Optional, Dict, Any
import logging
import re

logger = logging.getLogger(__name__)


class TemplateDocumentGenerator:
    """Generates documents by filling in Word template placeholders."""

    def generate_letter_from_template(
        self,
        template_path: Path,
        output_path: Path,
        fields: Dict[str, str]
    ) -> Path:
        """
        Generate a letter from a template by building the body content.
        Preserves header/footer (letterhead) from template.

        Args:
            template_path: Path to template .docx file with header/footer
            output_path: Where to save the generated document
            fields: Dict of field values (date, recipient_name, body, etc.)

        Returns:
            Path to generated document
        """
        try:
            # Load template (has header/footer with letterhead)
            doc = Document(str(template_path))

            # Clear existing body paragraphs
            for paragraph in doc.paragraphs:
                paragraph.clear()

            # Remove all existing paragraphs
            for _ in range(len(doc.paragraphs)):
                doc._element.body.remove(doc.paragraphs[0]._element)

            # Build the letter body
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

            # Recipient name and date on same line (name left, date right)
            first_line = doc.add_paragraph()
            # Add recipient name
            first_line.add_run(fields.get('recipient_name', ''))
            # Add tab and date on right
            tab_stops = first_line.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)
            first_line.add_run('\t' + fields.get('date', ''))
            doc.add_paragraph()  # blank line

            # Rest of recipient address
            if fields.get('recipient_company'):
                doc.add_paragraph(fields.get('recipient_company'))
            # Split address by newlines
            address = fields.get('recipient_address', '')
            for line in address.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
            doc.add_paragraph()  # blank line

            # Subject line (if provided)
            if fields.get('subject'):
                doc.add_paragraph(f"Re: {fields.get('subject')}")
                doc.add_paragraph()

            # Salutation
            doc.add_paragraph(fields.get('salutation', 'Dear Sir or Madam,'))
            doc.add_paragraph()

            # Body text (split by paragraphs)
            body = fields.get('body', '')
            for para_text in body.split('\n\n'):
                if para_text.strip():
                    # Handle single newlines within a paragraph
                    para_text = para_text.replace('\n', ' ')
                    doc.add_paragraph(para_text.strip())

            doc.add_paragraph()  # blank line

            # Closing (right-aligned / heavily indented)
            closing_para = doc.add_paragraph(fields.get('closing', 'Sincerely,'))
            closing_para.alignment = 2  # Right alignment
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()

            # Signature name (right-aligned / heavily indented)
            signature_para = doc.add_paragraph(fields.get('signature_name', ''))
            signature_para.alignment = 2  # Right alignment
            doc.add_paragraph()

            # Initials (left-aligned)
            doc.add_paragraph(fields.get('initials', ''))

            # Enclosures (if provided)
            if fields.get('enclosures'):
                enclosures = fields.get('enclosures')
                # Convert to string if it's a list
                if isinstance(enclosures, list):
                    enclosures = 'Enc: ' + ', '.join(enclosures)
                else:
                    enclosures = str(enclosures)
                    # Format as "Enc:" if not already formatted
                    if not enclosures.startswith('Enc:') and not enclosures.startswith('Enclosures:'):
                        enclosures = 'Enc: ' + enclosures
                doc.add_paragraph(enclosures)

            # Save generated document
            doc.save(str(output_path))

            logger.info(f"Generated letter: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error generating letter: {e}", exc_info=True)
            raise ValueError(f"Failed to generate letter: {str(e)}")

    def replace_placeholders_in_doc(
        self,
        template_path: Path,
        output_path: Path,
        placeholders: Dict[str, str]
    ) -> Path:
        """
        Replace placeholders in a Word document template.

        Args:
            template_path: Path to template .docx file
            output_path: Where to save the generated document
            placeholders: Dict of placeholder_name -> replacement_value

        Returns:
            Path to generated document
        """
        try:
            # Load template
            doc = Document(str(template_path))

            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_in_paragraph(paragraph, placeholders)

            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, placeholders)

            # Replace placeholders in headers/footers
            for section in doc.sections:
                # Header
                header = section.header
                for paragraph in header.paragraphs:
                    self._replace_in_paragraph(paragraph, placeholders)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                self._replace_in_paragraph(paragraph, placeholders)

                # Footer
                footer = section.footer
                for paragraph in footer.paragraphs:
                    self._replace_in_paragraph(paragraph, placeholders)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                self._replace_in_paragraph(paragraph, placeholders)

            # Save generated document
            doc.save(str(output_path))

            logger.info(f"Generated document: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error generating document: {e}", exc_info=True)
            raise ValueError(f"Failed to generate document: {str(e)}")

    def _replace_in_paragraph(self, paragraph, placeholders: Dict[str, str]):
        """Replace placeholders in a paragraph while preserving formatting."""
        # Get full text
        full_text = paragraph.text

        # Check if any placeholders exist
        has_placeholders = False
        for key in placeholders.keys():
            if f"{{{{{key}}}}}" in full_text:
                has_placeholders = True
                break

        if not has_placeholders:
            return

        # Replace placeholders in full text
        new_text = full_text
        for key, value in placeholders.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in new_text:
                # Convert None to empty string
                replacement = value if value is not None else ""
                new_text = new_text.replace(placeholder, replacement)

        # Clear paragraph and add new text (preserves first run's formatting)
        if paragraph.runs:
            # Keep first run's formatting
            first_run = paragraph.runs[0]
            paragraph.clear()
            new_run = paragraph.add_run(new_text)

            # Copy formatting from first run
            new_run.bold = first_run.bold
            new_run.italic = first_run.italic
            new_run.underline = first_run.underline
            new_run.font.name = first_run.font.name
            new_run.font.size = first_run.font.size
        else:
            paragraph.text = new_text

    def convert_to_pdf(self, docx_path: Path, pdf_path: Path) -> Path:
        """
        Convert Word document to PDF.

        Note: This requires docx2pdf or similar library.
        For now, just return the docx path.

        Args:
            docx_path: Path to .docx file
            pdf_path: Where to save PDF

        Returns:
            Path to PDF (or docx if conversion not available)
        """
        try:
            # Try to import docx2pdf
            from docx2pdf import convert
            convert(str(docx_path), str(pdf_path))
            logger.info(f"Converted to PDF: {pdf_path}")
            return pdf_path
        except ImportError:
            logger.warning("docx2pdf not available, returning Word document instead")
            return docx_path
        except Exception as e:
            logger.error(f"Error converting to PDF: {e}")
            return docx_path


# Singleton instance
_template_generator: Optional[TemplateDocumentGenerator] = None


def get_template_generator() -> TemplateDocumentGenerator:
    """Get or create the singleton TemplateDocumentGenerator instance."""
    global _template_generator
    if _template_generator is None:
        _template_generator = TemplateDocumentGenerator()
    return _template_generator
